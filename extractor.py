"""
extractor.py
------------
Handles text extraction from supported document types:
  - PDF  (.pdf)  via PyPDF2
  - DOCX (.docx) via python-docx
  - TXT  (.txt)  via built-in open()

All public functions raise ValueError for unsupported types and
IOError / specific library exceptions for read failures.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """
    Extract and return all text content from the given file.

    Args:
        file_path: Absolute or relative path to a .pdf, .docx, or .txt file.

    Returns:
        A single string containing all extracted text.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file extension is not supported.
        RuntimeError: If extraction fails for any other reason.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".txt":
        return _extract_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            "Supported types: .pdf, .docx, .txt"
        )


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        import PyPDF2  # local import so missing dep only fails for PDF files
    except ImportError as exc:
        raise RuntimeError(
            "PyPDF2 is not installed. Run: pip install PyPDF2"
        ) from exc

    pages: list[str] = []
    try:
        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            if reader.is_encrypted:
                logger.warning(
                    "PDF '%s' is encrypted. Attempting decryption with empty password.",
                    path.name,
                )
                reader.decrypt("")

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
                else:
                    logger.debug("Page %d of '%s' yielded no text.", page_num, path.name)

    except PyPDF2.errors.PdfReadError as exc:
        raise RuntimeError(f"Failed to read PDF '{path}': {exc}") from exc

    full_text = "\n\n".join(pages)
    if not full_text.strip():
        raise RuntimeError(
            f"No readable text found in PDF '{path}'. "
            "The file may be image-based (scanned) or corrupted."
        )

    logger.info("Extracted %d characters from PDF '%s'.", len(full_text), path.name)
    return full_text


def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document  # local import
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX '{path}': {exc}") from exc

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    full_text = "\n".join(paragraphs)
    if not full_text.strip():
        raise RuntimeError(
            f"No readable text found in DOCX '{path}'. The document may be empty."
        )

    logger.info("Extracted %d characters from DOCX '%s'.", len(full_text), path.name)
    return full_text


def _extract_txt(path: Path) -> str:
    """Extract text from a plain text file."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    last_error: Exception | None = None

    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as fh:
                full_text = fh.read()
            if not full_text.strip():
                raise RuntimeError(
                    f"The text file '{path}' is empty."
                )
            logger.info(
                "Extracted %d characters from TXT '%s' (encoding: %s).",
                len(full_text),
                path.name,
                encoding,
            )
            return full_text
        except UnicodeDecodeError as exc:
            last_error = exc
            logger.debug(
                "Encoding '%s' failed for '%s', trying next.", encoding, path.name
            )

    raise RuntimeError(
        f"Could not decode '{path}' with any supported encoding. "
        f"Last error: {last_error}"
    )
